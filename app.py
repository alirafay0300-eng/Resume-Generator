from flask import Flask, render_template, request, make_response, send_file, session, redirect, url_for, flash
from weasyprint import HTML, CSS
from docx import Document
import io
import os
from werkzeug.utils import secure_filename
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
#this is a change 
# ==============================
# App Config
#this is a change
# ==============================
app = Flask(__name__)
app.secret_key = "your_secret_key"

# Database
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///payments.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)
migrate = Migrate(app, db)


# Upload folder
UPLOAD_FOLDER = 'static/uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Admin
ADMIN_PASSWORD = "your_password"

# ==============================
# Database Models
# ==============================
class Payment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    txn_id = db.Column(db.String(100), nullable=False)
    amount = db.Column(db.Float, nullable=False)
    screenshot = db.Column(db.String(200))  # filename
    status = db.Column(db.String(20), default='Pending')  # Pending / Verified
    payment_status = db.Column(db.String(20), default="pending")  # can be 'pending', 'approved', 'rejected'


# ✅ Create tables inside application context
with app.app_context():
    db.create_all()


# ==============================
# Resume storage
# ==============================
user_resume_data = {}

# ==============================
# Routes
# ==============================

@app.route('/')
def form():
    return render_template('form.html')
@app.route('/submit', methods=['POST'])
@app.route('/submit', methods=['POST'])
def submit():
    # Collect data (education, experience, projects, skills)
    educations = [
        (edu, uni)
        for edu, uni in zip(request.form.getlist('education[]'), request.form.getlist('university[]'))
        if edu.strip() or uni.strip()
    ]

    experiences = []
    jobs = request.form.getlist('job[]')
    companies = request.form.getlist('company[]')
    starts = request.form.getlist('start[]')
    ends = request.form.getlist('end[]')
    descs = request.form.getlist('desc[]')
    continued_flags = request.form.getlist('continued[]')

    for i, job in enumerate(jobs):
        if not job.strip():
            continue
        company = companies[i] if i < len(companies) else ""
        start = starts[i] if i < len(starts) else ""
        end = ends[i] if i < len(ends) else ""
        desc = descs[i] if i < len(descs) else ""
        is_continued = continued_flags[i] == 'on' if i < len(continued_flags) else False
        period = f"{start} to {end}" if not is_continued else f"{start} to Present"
        experiences.append({'job': job, 'company': company, 'period': period, 'desc': desc})

    projects = [
        (t, d)
        for t, d in zip(request.form.getlist('project_title[]'), request.form.getlist('project_detail[]'))
        if t.strip() or d.strip()
    ]

    skills = [s for s in request.form.getlist('skills[]') if s.strip()]

    # Store resume data globally
    global user_resume_data
    user_resume_data = {
        'name': request.form.get('name', ''),
        'address': request.form.get('address', ''),
        'phone': request.form.get('phone', ''),
        'email': request.form.get('email', ''),
        'educations': educations,
        'experiences': experiences,
        'projects': projects,
        'skills': skills
    }


    session['user_name'] = request.form.get('name', '')
    return render_template('resume.html', **user_resume_data)


@app.route('/download_docx')
def download_docx():
    global user_resume_data
    if not user_resume_data:
        return "Please generate your resume first.", 400

    name = session.get('user_name')
    if not name:
        flash("Please submit your resume first.")
        return redirect(url_for('form'))

    payment = Payment.query.filter_by(name=name).first()  # Any payment

    if not payment:
        # No payment submitted yet → ask to pay
        return redirect(url_for('payment_page'))

    if payment.status != "Verified":
        # Payment submitted but not verified
        return render_template('payment_wait.html')

    # Payment verified → generate DOCX
    doc = Document()
    doc.add_heading(user_resume_data.get('name', ''), 0)
    doc.add_paragraph(f"{user_resume_data.get('address', '')} | {user_resume_data.get('phone', '')} | {user_resume_data.get('email', '')}")

    # Education
    doc.add_heading('Education', level=1)
    for edu, uni in user_resume_data.get('educations', []):
        doc.add_paragraph(f"{edu} – {uni}")

    # Experience
    doc.add_heading('Experience', level=1)
    for exp in user_resume_data.get('experiences', []):
        if exp['job']:
            p = doc.add_paragraph()
            p.add_run(f"{exp['job']} – {exp['company']}").bold = True
            p.add_run(f"    {exp['period']}")
            if exp['desc']:
                for line in exp['desc'].splitlines():
                    doc.add_paragraph(f"• {line}", style='ListBullet')

    # Projects
    doc.add_heading('Projects', level=1)
    for title, detail in user_resume_data.get('projects', []):
        doc.add_paragraph(f"{title}: {detail}")

    # Skills
    doc.add_heading('Technical Skills', level=1)
    skills = user_resume_data.get('skills', [])
    if skills:
        doc.add_paragraph(', '.join(skills))

    # Send file
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return send_file(
        f,
        as_attachment=True,
        download_name='resume.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/download_pdf')
def download_pdf():
    global user_resume_data

    if not user_resume_data:
        return "Please generate your resume first.", 400

    # ✅ Check payment approval
    # You can store the user's name or ID in session when they submit payment
    name = session.get('user_name')
    payment = Payment.query.filter_by(name=name).first()  # Any payment

    if not payment:
        # No payment submitted yet → ask to pay
        return redirect(url_for('payment_page'))

    if payment.status != "Verified":
        # Payment submitted but not verified
        return render_template('payment_wait.html')

    # Generate PDF
    rendered = render_template('resume.html', **user_resume_data)
    css_path = os.path.join('static', 'main.css')
    pdf = HTML(string=rendered, base_url=request.root_path).write_pdf(
        stylesheets=[CSS(css_path)]
    )

    response = make_response(pdf)
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = 'attachment; filename=resume.pdf'
    return response

# ==============================
# Admin
# ==============================

@app.route('/admin_login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        if request.form['password'] == ADMIN_PASSWORD:
            session['admin'] = True
            return redirect(url_for('admin_dashboard'))
        else:
            flash("Wrong password!")
    return render_template('admin_login.html')

@app.route('/admin')
def admin_dashboard():
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    payments = Payment.query.all()
    return render_template('admin_dashboard.html', payments=payments)

@app.route('/update_payment/<int:payment_id>/<string:action>', methods=['GET', 'POST'])
def update_payment(payment_id, action):
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    payment = Payment.query.get_or_404(payment_id)
    if action == "verify":
        payment.status = "Verified"
        payment.payment_status = "approved"
    elif action == "reject":
        payment.status = "Rejected"
        payment.payment_status = "rejected"
    db.session.commit()
    flash("Payment updated successfully!")
    return redirect(url_for('admin_dashboard'))


@app.route('/payment_submit', methods=['POST''GET'])
def payment_submit():
    name = request.form.get('name')
    txn_id = request.form.get('txn_id')
    amount = request.form.get('amount')
    screenshot = request.files.get('screenshot')
    filename = None
    if screenshot:
        filename = secure_filename(screenshot.filename)
        screenshot.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    
    payment = Payment(name=name, txn_id=txn_id, amount=float(amount), screenshot=filename)
    db.session.add(payment)
    db.session.commit()

    # Optional: session can indicate payment submitted
    session['payment_submitted'] = True

    # Show waiting message
    return render_template('payment_wait.html')

@app.route('/payment')
def payment_page():
    if 'user_name' not in session:
        flash("Please submit your resume first.")
        return redirect(url_for('form'))
    return render_template('payment.html', name=session['user_name'])

@app.route("/approve/<int:payment_id>")
def approve_payment(payment_id):
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    payment = Payment.query.get_or_404(payment_id)
    payment.status = "Verified"
    payment.payment_status = "approved"
    db.session.commit()
    flash("Payment approved!")
    return redirect(url_for("admin_dashboard"))



@app.route("/waiting")
def waiting():
    return render_template("waiting.html")
@app.route("/thank_you/<int:user_id>")


@app.route("/check_status/<int:payment_id>")
def check_status(payment_id):
    payment = Payment.query.get_or_404(payment_id)
    if payment.payment_status == "approved":
        return redirect(url_for("thank_you", payment_id=payment.id))
    else:
        return render_template("waiting.html")

@app.route("/thank_you/<int:payment_id>")
def thank_you(payment_id):
    payment = Payment.query.get_or_404(payment_id)
    return render_template("thank_you.html", payment=payment)

@app.route("/download_pdf_user/<int:payment_id>")
def download_pdf_user(payment_id):
    payment = Payment.query.get_or_404(payment_id)
    if payment.payment_status != "approved":
        flash("Payment not approved yet!")
        return redirect(url_for("waiting"))
    return send_file("path/to/file.pdf", as_attachment=True)

# ==============================
# Run App
# ==============================
if __name__ == '__main__':
    app.run(debug=True)
